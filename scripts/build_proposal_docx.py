"""Build a minimalistic .docx of the NS-RAM research proposal."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path("/home/ikaros/Documents/claude_hive/AMD_gfx1151_energy/docs/NSRAM_Research_Proposal_2026-04-30.docx")

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

# Default body style
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)


def H1(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(16)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    return p


def H2(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    return p


def H3(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(10.5)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    return p


def P(text, italic=False, size=10.5, space_after=4):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text).font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(2)
    return p


def hr():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:color"), "BBBBBB")
    pBdr.append(bottom)
    pPr.append(pBdr)


# --- Title block --- #
H1("NS-RAM Co-Design Collaboration")
P("Research Proposal · 30 April 2026", italic=True, size=10)
P("Eric Bergvall · Robert Luciani  →  Prof. Mario Lanza · Dr. Sebastian Pazos",
  size=10, space_after=10)
hr()

# --- Executive summary --- #
H2("Executive summary")
P("Sebastian's NS-RAM 2T cell can act as synapse, neuron soma, or "
  "integrate-and-fire spiker depending on its bias regime — a property no "
  "memristor or single-element device offers. Realising this in silicon for "
  "a useful workload requires a tight loop between measured device, "
  "fast-and-accurate simulator, and AI architecture. That loop is currently "
  "broken: device measurement and fab cycles take months; algorithmic "
  "exploration takes minutes. We close it.")
P("Eric (differentiable PyTorch port of BSIM4 + parasitic NPN, calibrated "
  "on Sebastian's 33 measured curves) and Robert (Julia/MTK DEQ inner solver, "
  "71 000 root evaluations per second on a single GB10) propose to serve as "
  "the software bridge between Mario's foundry strategy, Sebastian's silicon, "
  "and the AI workloads the chip must run. Two independently built simulators "
  "that agree with each other and with measured silicon let the hardware team "
  "test chip variants at the speed of software, and let the algorithm team "
  "target physically realistic parameter values from day one.")
P("We propose a 12-month engagement with three concrete deliverables, the "
  "first ready in 8 weeks, directly feeding Mario's TSMC test-vehicle "
  "planning and the 6 May NRF submission.")

# --- 1. Opportunity --- #
H2("1.  The opportunity has a deadline")
P("NS-RAM's strategic window is 18–24 months, not seven years. Frontier AI "
  "is moving from frozen feed-forward inference to adaptive, online, "
  "variable-compute systems. The 2T cell's capacity to dynamically reassign "
  "roles (synapse ↔ soma ↔ IF spiker through VG2) maps directly onto this "
  "direction; a competing array of memristors or analogue MAC tiles cannot. "
  "The longer the gap between device characterisation and a published "
  "algorithmic demonstration, the more this advantage erodes. Our role is "
  "to compress the loop.")

# --- 2. USPs --- #
H2("2.  Two USPs we bring")

H3("USP 1 — Speed: 600 000× headroom over today's literature tooling")
table = doc.add_table(rows=7, cols=2)
table.autofit = True
rows = [
    ("Stack", "Throughput (cell evals / s)"),
    ("ngspice + BSIM4 (literature)", "50"),
    ("Canonical Newton, 20-thread CPU", "22 000"),
    ("DEQ inner solver, GB10 (today)", "71 000"),
    ("GPU-only ceiling on GB10", "441 000"),
    ("GB10 + custom kernels (Q3 2026)", "3 000 000"),
    ("B200 + full optimisation (2027)", "30 000 000"),
]
for i, (a, b) in enumerate(rows):
    c0, c1 = table.rows[i].cells
    c0.text = a
    c1.text = b
    for c in (c0, c1):
        for p in c.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9.5)
                if i == 0:
                    r.bold = True
P("Sebastian asks 'what if we widen the M2 channel by 20 %?' — Robert returns "
  "the full I-V family plus three downstream network-task results in the "
  "same meeting. Today that question takes a fab cycle.",
  size=9.5, italic=True)

H3("USP 2 — Accuracy through 2T-cell-aware co-design")
P("The 2T NS-RAM cell is not a generic transistor pair. Its M1 + M2 + "
  "parasitic NPN topology with a floating P-body creates regime-switching "
  "behaviour that BSIM4 alone cannot describe. We have built the cell-level "
  "wrapper (Newton-Raphson on Vsint and Vb; pseudo-arclength continuation "
  "through the snapback fold; differentiable end-to-end) that turns BSIM4 "
  "into a faithful NS-RAM simulator. Cross-validated:")
bullet("Vth assembly: ±0.1 mV vs ngspice across 360 bias points")
bullet("Id (saturation): ±2 % on M2 (long), ±10–15 % on M1 (short)")
bullet("All BSIM4 sub-formulas line-by-line audited against Berkeley v4.8.3")
bullet("Pseudo-arclength solver: 100 % Newton convergence across the snapback knee")
P("Two independent simulators (Eric's PyTorch + Robert's Julia/MTK) "
  "cross-check each other; a result counts only when both stacks agree.",
  size=9.5, italic=True)

# --- 3. Deliverables --- #
H2("3.  What we deliver — concrete, dated, foundry-relevant")

H3("D1 — Calibrated NS-RAM digital twin · 8 weeks  (feeds NRF + TSMC vehicle)")
bullet("BSIM4 + NPN port fitted to Sebastian's 33 curves, anchored to his "
       "extracted VTH0 / BETA0 / ETA0 / NFACTOR trajectories.")
bullet("Robert's MTK stack reproduces the fit on a held-out grid "
       "(≤ 2 % per-curve log-RMSE deviation between the two simulators).")
bullet("Energy figures derived directly from Sebastian's published numbers "
       "(~21 fJ/cycle, ~6.7 fJ/spike, ~46 µm² per cell).")
bullet("Output: nsram_cell package v0.13 + a one-page energy/accuracy summary "
       "for Mario's NRF portal submission.")

H3("D2 — Architecture exploration platform · 6 months")
bullet("Standardised benchmark suite (Hopfield, memory capacity, NARMA-10, "
       "7-class waveform, temporal-XOR) running at 71 k+ evaluations/s.")
bullet("Bilevel exploration loop: Eric's autograd descends on cell parameters "
       "for fixed architecture; Robert's MTK sweeps architectures for fixed "
       "parameters; cross-validation gate prevents simulator artefacts.")
bullet("Sebastian uses the platform to explore physically reachable bias "
       "regions; Mario uses the resulting design-rule table to size the next "
       "tape-out.")
bullet("Output: benchmark-suite repo + design-rule sheet for the foundry.")

H3("D3 — Headline demonstration: cognitive core on NS-RAM · 12 months")
bullet("Distil a 1–3 B-parameter reasoning core from a frontier model "
       "(LASER-style structured pruning, reasoning-calibrated).")
bullet("Convert to looped-transformer form with adaptive-halt value head — "
       "variable-compute by construction.")
bullet("Co-design the NS-RAM array geometry for the access patterns the "
       "algorithm needs; tape out a small array; characterise.")
bullet("Headline claim: a single device that is synapse, soma and IF spiker, "
       "carrying a reasoning model that thinks as long as it needs to — "
       "uniquely NS-RAM.")
bullet("Output: preprint + 2-page foundry technical sheet for the next "
       "tape-out cell library.")
P("Each deliverable is independently useful. If D3 slips, D1 + D2 are still "
  "wins. If D2 slips, D1 still feeds Mario's NRF page.",
  size=9.5, italic=True)

# --- 4. Three questions --- #
H2("4.  Three questions Mario will be asked, with answers")

H3("$50 M for seven years — but AI moves in 12 months. What's deliverable in 12?")
P("D1 in 8 weeks. D2 in 6 months. A first algorithmic demonstration on the "
  "calibrated twin in 9–12 months. The seven-year frame buys the array "
  "tape-out and the integrated demonstrator; the 12-month frame buys the "
  "credibility to commit to it.")

H3("Universities are full of hype. What concretely scales?")
P("The single device with three regimes does. Measured energy is ~6.7 fJ "
  "per spike at ~46 µm². Anything that maps to spike-driven sparse "
  "computation maps onto it. We commit to publishing the cross-validated "
  "energy/accuracy table, with all simulator code open, in the first 8 weeks.")

H3("Why this chip and not a 1 W signal-processing accelerator?")
P("A 1 W signal-processing accelerator runs frozen models well. It cannot "
  "run adaptive, variable-compute, or online-learning algorithms — those "
  "need a substrate where weights and dynamics are learnable in place. "
  "NS-RAM is one of the very few candidate substrates. Our role is to "
  "prove this with a demonstration, not a slide.")

# --- 5. Open questions --- #
H2("5.  Key open questions")
H3("To Sebastian")
bullet("Raw transient (Vd(t), Id(t), VG2(t)) traces behind slide 25 — "
       "available? Needed for body-charge τ-calibration in D1.")
bullet("Extracted parameter curves (slide 24) — available as CSV? "
       "We would use them as fitting anchors directly.")
bullet("~9 missing sweeps in the high-VG2 + high-Is corner where LDE "
       "physics dominates — feasible to add?")
H3("To Mario")
bullet("Target array size for the next tape-out — sets D2's scaling sweep "
       "(N = 64 → 512 → 2048 cells).")
bullet("Power-band positioning — slide 26 places NS-RAM in the 10–100 mW "
       "gateway tier. Confirmed?")
bullet("NRF one-pager (6 May) — do you want us to draft the technical third "
       "by 3 May, or fill yours?")
bullet("Headline: meta-plasticity (cell-role self-allocation) versus "
       "variable-compute reasoning (adaptive-halt cognitive core) — "
       "which resonates more with Newmorphic's positioning?")

# --- 6. Engagement --- #
H2("6.  Proposed engagement structure")
P("This is a collaboration, not a subcontract. Structure:")
bullet("Weekly 30-minute sync (the four of us), starting now.")
bullet("Quarterly review with one shared written status (~3 pages).")
bullet("Joint authorship on outputs that depend on both sides.")
bullet("Open simulator code with NDA-equivalent protection for Sebastian's "
       "unpublished measurements until publication.")
bullet("IP: simulator and algorithmic IP with Eric/Robert side; cell, "
       "process, tape-out IP with Mario/Sebastian side; joint IP on "
       "co-designed cell + algorithm artefacts. Refined in the agreement.")
P("Memorandum of understanding within two weeks; full collaboration "
  "agreement within six.")

# --- 7. Why now --- #
H2("7.  Why now")
P("Mario has the foundry path and the funding window. Sebastian has the "
  "device and the measurements. Robert and Eric have the speed and the "
  "accuracy. Each of the four pieces is mature individually; none of them "
  "delivers on its own. The proposal is to wire them together for the next "
  "12 months and see what we can build.")

hr()
P("Eric Bergvall  ·  Robert Luciani  ·  30 April 2026", italic=True, size=9)

doc.save(str(OUT))
print(f"Saved {OUT}")
