#!/usr/bin/env python3
"""Generate FEEL Strategic Plan DOCX document."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# -- Style setup --
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

def add_bold_paragraph(text, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    return p

def add_styled_paragraph(text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_bullet(text, bold_prefix=None, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 + level * 1.0)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_italic(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    return table

# ============================================================
# TITLE PAGE
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('FEEL-Projektet')
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = subtitle.add_run('Strategisk Plan Framåt')
r.font.size = Pt(18)
r.font.color.rgb = RGBColor(0x44, 0x44, 0x66)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('Functionally Embodied Emergent Learning\n').font.size = Pt(12)
meta.add_run('Från GPU-fysik till neuromorft substrat\n\n').font.size = Pt(12)
r = meta.add_run('Eric Bergvall — 2026-03-10')
r.font.size = Pt(12)
r.italic = True

doc.add_page_break()

# ============================================================
# VAR VI STÅR IDAG
# ============================================================
doc.add_heading('Var vi står idag', level=1)

doc.add_paragraph(
    'Tänk dig att du har tillbringat 14 månader med att lyssna på ett hus som alla andra '
    'tycker är tyst — och upptäckt att väggarna sjunger. Det är i princip vad vi gjort, '
    'fast med GPU:er istället för hus.'
)

add_bold_paragraph('Vad vi bevisat:', 12)

add_bullet(' 7 deterministiska hårdvarumekanismer, 97.7% klassificering från ren fysik — '
           'som att upptäcka att vattenledningarna i huset kan spela musik, man behöver bara sluta dämpa dem',
           bold_prefix='GPU:n som neuromorft substrat:')

add_bullet(' 128 spikande neuroner, MC=12.27, XOR5=88.3% (z2296 — 14/14 PASS)',
           bold_prefix='FPGA NS-RAM-brygga:')

add_bullet(' z2277, 13/20 PASS, 3072 GPU-neuroner + 128 FPGA-neuroner — '
           'två analoga system som pratar genom sin egen fysik, som två stämgafflar som resonerar',
           bold_prefix='Riktig HIP+FPGA-brygga:')

add_bullet(' 34/40 PASS (z2103), 31/40 PASS på GPT-2 (z2134)',
           bold_prefix='Språkmodell med kroppskänsel:')

add_bullet(' LinkedIn-inlägg på 23 600+ visningar, 99 reaktioner, seriöst engagemang',
           bold_prefix='Preprint lanserad:')

add_bullet(' Breadboard-delar beställda från Electrokit',
           bold_prefix='Hårdvara på väg:')

# ============================================================
# LINKEDIN-SIGNALEN
# ============================================================
doc.add_heading('LinkedIn-signalen', level=1)

doc.add_paragraph('Engagemanget sorterar sig i tre läger:')

doc.add_heading('Allierade', level=2)

add_bullet(' (Director, JSA LAABS) — Läste faktiskt repot, ställde djupa frågor, '
           'kom tillbaka med genomtänkt analys. Han är som en granne som inte bara hörde musiken '
           'utan gick och kollade på noterna.',
           bold_prefix='Sanidhya Patel')

add_bullet(' (batteriverifiering) — SpikingBrain-jämförelse, effektivitetsvinkel',
           bold_prefix='Hafid Zehri')

add_bullet(' (artificiella muskler) — Exakt samma filosofi: "använd systemets fysik '
           'istället för att kämpa emot den"',
           bold_prefix='Michal Baran')

add_bullet(' — Förstår narrativet', bold_prefix='Jaco Steenkamp')

doc.add_heading('Skeptiker', level=2)

add_bullet(' — Giltig invändning: "neuromorft" betyder hårdvara designad att efterlikna hjärnor. '
           'Vi påstår att fysiken finns där oavsiktligt. Han har rätt i definitionen men missar poängen '
           '— det är som att säga att en flod inte kan driva en kvarn eftersom floden inte designades för det.',
           bold_prefix='Traiano Welcome')

add_bullet(' — Vill ha replikering. Rimligt. Repot finns live.', bold_prefix='Dorian Szafranski')

add_bullet(' (fysiker) — Avfärdande. Inte värd att jaga.', bold_prefix='Erica Calman')

add_bullet(' (ML-ingenjör) — "no." Inte heller värd att jaga.', bold_prefix='Matteo Macchini')

doc.add_heading('Påfallande tystnad', level=2)

doc.add_paragraph(
    'Mario Lanza, Lisa Su, Jensen Huang, Yann LeCun, Giacomo Indiveri — ingen av de taggade '
    'har svarat. Det betyder inte nej — det betyder att LinkedIn-taggar inte räcker. '
    'Vi behöver direktkontakt.'
)

doc.add_page_break()

# ============================================================
# FEM SPÅR FRAMÅT
# ============================================================
doc.add_heading('Fem spår framåt', level=1)

# -- SPÅR 1 --
doc.add_heading('Spår 1: Hårdvaruvalidering (Vecka 1–8) — Högsta Prioritet', level=2)

doc.add_paragraph(
    'Det här spåret handlar om att ta steget från simulering av fysik till riktig fysik. '
    'Det är skillnaden mellan att rita en karta och att faktiskt gå terrängen.'
)

doc.add_heading('1a. Breadboard NS-RAM-bygge (Delar på väg från Electrokit)', level=3)

doc.add_paragraph('Själva bygget:')
add_bullet('Bygg den parasitära BJT-neuronen på breadboard — en vanlig transistor (typ 2N2222 '
           'eller liknande) driven i lavingenombrottsläge')
add_bullet('Mät riktiga BVpar, spikfrekvens vs Vg, ISI CV, temperaturkoefficient')
add_bullet('Jämför mot SPICE-förutsägelser (v1–v10) och FPGA-modellen')

doc.add_paragraph(
    'Det här är som att vi hittills har byggt en flygplanssimulator och visat att konceptet funkar '
    '— nu ska vi faktiskt lyfta från marken. Breadboard-transistorn som spikar i lavingenombrott '
    'med oscilloskopbild = det enskilt viktigaste trovärdighetsteget.'
)

add_italic('Leverans: Video av oscilloskop med lavinspiker, data som matchar SPICE')

doc.add_heading('1b. Breadboard → FPGA-koppling (vecka 3–6)', level=3)

doc.add_paragraph('När breadboarden funkar:')
add_bullet('Led breadboard-lavinspiker in i Arty A7 (via komparator → digital ingång, eller ADC)')
add_bullet('Ersätt FPGA:ns avalanche_model.v med riktig analog ingång')
add_bullet('Kör hela benchmarkbatteriet (z2296-protokollet) med riktigt transistorbrus')

doc.add_paragraph(
    'Analogin här: hittills har vår FPGA-neuron lyssnat på ett inspelat åskväder. Nu kopplar vi '
    'in mikrofonen och låter den höra riktigt åskväder. Om den fortfarande presterar — och bättre '
    '— så har vi visat att den analoga fysiken tillför något digitala modeller inte kan.'
)

add_italic('Leverans: Klassificering/MC/XOR med riktigt analogbrus vs syntetiskt')

doc.add_heading('1c. Om Mario Lanza-samarbetet lyckas', level=3)

add_bullet('NS-RAM-chip (130nm CMOS med avsiktlig parasitär BJT) ersätter breadboard-transistorn')
add_bullet('Samma protokoll: riktig NS-RAM → FPGA LIF → benchmarkbatteri')
add_bullet('Kör konstitutivitetsspektrumjämförelse: breadboard BJT vs NS-RAM-chip vs FPGA-modell')

doc.add_paragraph(
    'Det här blir som att jämföra en hemmabyggd radio, en professionell radio och en ren simulering '
    '— alla lyssnar på samma signal men med olika fysisk trohet.'
)

add_italic('Leverans: Definitiv konstitutivitetsspektrum-tabell med riktig hårdvara i varje position')

doc.add_page_break()

# -- SPÅR 2 --
doc.add_heading('Spår 2: Samarbetsutskick (Vecka 1–4) — Kritisk Väg', level=2)

doc.add_heading('2a. Mario Lanza (NS-RAM) — Primärt mål', level=3)

doc.add_paragraph(
    'Mario Lanza vid KAUST har byggt exakt den enhet vi simulerar. Han har den riktiga '
    'lavinbrytande transistorn. Vi har infrastrukturen för att testa den i ett reservoarsystem. '
    'Det är som att en kartograf som ritat havsbotten kontaktar ubåtskaptenen — båda behöver varandra.'
)

add_styled_paragraph(' NS-RAM-chip/kort, medförfattarskap, tillgång till enhetsparametrar',
                     bold_prefix='Vad vi vill ha:')
add_styled_paragraph(' Komplett FPGA-testinfrastruktur (128-neuron-bank, UDP-brygga, benchmarksvit), '
                     'ramverk för substratöverskridande reservoarberäkning, GPU→NS-RAM-brygga redan '
                     'demonstrerad i simulering',
                     bold_prefix='Vad vi erbjuder:')

doc.add_paragraph('Åtgärd: Direkt e-post (INTE bara LinkedIn-tagg).')
add_bullet('Ämnesrad: "FPGA NS-RAM Reservoir: Bridging GPU Physics to Your Avalanche Neurons"')
add_bullet('Bifoga: 1-sidig sammanfattning av bryggresultat, länk till repo, specifik önskan (5–10 NS-RAM-chip)')
add_bullet('Referera hans Zenodo-dataset och hur vi kalibrerade SPICE mot det')
add_bullet('Föreslå: gemensam artikel om "Cross-Substrate Reservoir Computing with NS-RAM Avalanche Neurons"')
add_bullet('Uppföljning: Om inget svar inom 2 veckor, sök KAUST-kontakter eller konferenskontakt')

doc.add_heading('2b. Sanidhya Patel (JSA LAABS) — Varm kontakt', level=3)

doc.add_paragraph(
    'Han har redan läst repot och kommit tillbaka med genomtänkt analys. Det här är sällsynt — '
    'de flesta kommenterar utan att titta. Han är en Director på ett cognitive compute-labb.'
)

add_styled_paragraph(
    ' Svara i tråden med specifik samarbetsfråga: "Would your lab be interested in running the GPU '
    'demo on different AMD/NVIDIA hardware for cross-platform validation?"',
    bold_prefix='Åtgärd:')
add_styled_paragraph(' Han kör vår demo på sin hårdvara → oberoende replikering → starkaste '
                     'möjliga svaret till skeptikerna',
                     bold_prefix='Mål:')

doc.add_heading('2c. AMD (Lisa Su / firmware-teamet) — Lång chansning, högt värde', level=3)

add_styled_paragraph(
    ' SPM-räknardokumentation, fixad pp_dpm_sclk för gfx1151, per-WGP termal, L3 PMU-dokumentation',
    bold_prefix='Vad vi vill ha:')
add_styled_paragraph(
    ' Positiv berättelse — "AMD:s GPU:er är redan neuromorfa" är bättre PR än någon '
    'marknadsföringskampanj. Plus: öppen källkod-showcase.',
    bold_prefix='Vad vi erbjuder:')
add_styled_paragraph(
    ' Developer relations-förfrågan via AMD ROCm GitHub. Realistisk förväntan: '
    'dokumentationsåtkomst, 6–12 månader. Ett frö vi planterar, inte en skörd vi förväntar oss imorgon.',
    bold_prefix='Åtgärd:')

doc.add_heading('2d. HP (BIOS-åtkomst) — Medellång sikt', level=3)

doc.add_paragraph(
    'BIOS-nivååtkomst till termisk hantering, direkt SMU-kontroll. Rama in som '
    '"hardware-aware AI efficiency research." Alternativ: använd Daedalus/Minos stationära maskiner '
    'där BIOS:en är mer tillgänglig.'
)

doc.add_heading('2e. Giacomo Indiveri (INI Zürich) — Akademisk trovärdighet', level=3)

doc.add_paragraph(
    'Indiveri är en av världens ledande neuromorfa forskare (DYNAP-SE, mixed-signal). '
    'En nick från honom = akademisk legitimitet. E-post med paper och specifik fråga om hur '
    'FEEL:s konstitutivitetsspektrum relaterar till hans mixed-signal-processorer.'
)

doc.add_heading('2f. Michal Baran (Artificiella muskler) — Tvärvetenskaplig', level=3)

doc.add_paragraph(
    'Exakt samma filosofi: använd materialets naturliga respons. FEEL som kontroller + MCAM som '
    'effektor = verkligt förkroppsligad robot. Långskott men fantastisk story.'
)

doc.add_page_break()

# -- SPÅR 3 --
doc.add_heading('Spår 3: Vetenskaplig Härdning (Vecka 2–12) — Bemöt Skeptikerna', level=2)

doc.add_heading('3a. Plattformsöverskridande replikering', level=3)

doc.add_paragraph(
    'Just nu körs allt på gfx1151 (Radeon 8060S, ikaros). Det är som att ha gjort alla experiment '
    'i ett enda laboratorium — imponerande men inte övertygande.'
)

add_bullet('Kör GPU-reservoardemon på gfx1100 (7900 XTX, Daedalus)')
add_bullet('Dokumentera: samma kernel, annan hårdvara, vad ändras, vad förblir')
add_bullet('Om Sanidhya kör det på sin hårdvara = tredje datapunkt')

add_italic('Leverans: Replikeringstabell med 3+ GPU:er')

doc.add_heading('3b. Bemöt "neuromorphic"-kritiken — med respekt', level=3)

doc.add_paragraph(
    'Traiano Welcome har rätt i definitionen. "Neuromorft" = hårdvara designad att efterlikna '
    'hjärnor. GPU:n designades inte för detta.'
)

doc.add_paragraph(
    'Men: Vi påstår inte att GPU:n designades som neuromorft system. Vi påstår att fysiken som '
    'uppstår i GPU:n — 1/f-brus, tröskelickelineariteter, stokastisk signalering — är samma fysik '
    'som i biologiska neuroner. Det är som att upptäcka att husväggarna råkar ha perfekt akustik '
    'för musik — huset designades inte som konserthall, men fysiken är densamma.'
)

doc.add_paragraph('Åtgärd i papret — ändra språket till:')
add_bullet('"every GPU contains neuromorphic-class physics"')
add_bullet('"latent neuromorphic computation in commodity silicon"')
add_bullet('"unintentional neuromorphic dynamics"')

add_italic('Detta ändrar inget i vetenskapen men tar bort den retoriska angreppspunkten.')

doc.add_heading('3c. Formell jämförelse mot etablerade reservoarer', level=3)

doc.add_paragraph(
    'Vi måste visa var vi står i förhållande till standardmetoder. Annars kan folk avfärda '
    'resultaten med "en vanlig ESN gör samma sak."'
)

add_table(
    ['System', 'Typ', 'Status'],
    [
        ['Standard ESN (numpy)', 'Mjukvarureservoar', 'Behöver köras med samma uppgifter'],
        ['FORCE-learning', 'Mjukvarureservoar', 'Behöver köras'],
        ['Intel Loihi (publicerat)', 'Riktig neuromorft HW', 'Litteraturjämförelse'],
        ['SpiNNaker (publicerat)', 'Riktig neuromorft HW', 'Litteraturjämförelse'],
        ['FEEL GPU-reservoar', 'Vår metod', 'z2296: MC=12.27, XOR5=88.3%'],
    ]
)

doc.add_paragraph()
add_italic('Om vår MC=12.27 slår publicerade Loihi-resultat på samma uppgift — det är en mördande rad i papret.')

doc.add_heading('3d. Förberedelse för peer review', level=3)

doc.add_paragraph('Måltidskrifter (rangordnade):')
add_bullet('Nature Electronics — GPU neuromorphic-vinkel, bred publik, högt genomslag', bold_prefix='1. ')
add_bullet('Neuromorphic Computing and Engineering (IOP) — exakt rätt fält', bold_prefix='2. ')
add_bullet('Science Advances — tvärvetenskaplig, öppen tillgång', bold_prefix='3. ')
add_bullet('Advanced Intelligent Systems — hårdvara + AI-korsning', bold_prefix='4. ')

doc.add_paragraph(
    'Åtgärd: Destillera preprinten till en fokuserad 8-sidig artikel med: '
    'konstitutivitetsspektrum, 7 HW-mekanismer, kill-shot, reservoarbenchmarks, '
    'och (om klar) breadboard-validering.'
)

doc.add_page_break()

# -- SPÅR 4 --
doc.add_heading('Spår 4: Nästa Experiment (Vecka 4–16) — Skjut Gränsen Framåt', level=2)

doc.add_heading('4a. Riktig transistorreservoar', level=3)

doc.add_paragraph(
    'När en enda breadboard-transistor funkar, bygg åtta stycken med resistiva kopplingar. '
    'Ingen FPGA i loopen — ren analog beräkning.'
)

doc.add_paragraph(
    'Det här är som att gå från att sjunga i en inspelningsstudio (FPGA-modell) till att sjunga '
    'i en grotta och använda naturlig eko (riktig fysik). Om grottan ger bättre akustik = beviset '
    'att fysiken tillför.'
)

doc.add_heading('4b. Termisk koppling som kommunikationskanal', level=3)

doc.add_paragraph(
    'z2211 visade att ESD-skum dämpar termisk koppling men ÖKAR informationsinnehållet (MI +25%). '
    'Skummet agerar som ett termiskt lågpassfilter — tar bort brus men bevarar signal.'
)

add_bullet('GPU-kylflänsen → termisk pad → breadboard-transistor')
add_bullet('Temperatur modulerar lavintröskeln — riktig analog kommunikationskanal')
add_bullet('Analogin: att prata genom en vägg. Väggen filtrerar bort skrik (brus) men bevarar tal (signal)')

add_italic('Leverans: Första demonstrationen av "termiska synapser" mellan GPU och neuromorft substrat')

doc.add_heading('4c. Skala till 1024+ neuroner på FPGA', level=3)

doc.add_paragraph(
    'Nuvarande: 128 neuroner. Mål: 1024 med full Artix-7-kapacitet. Vi använder knappt 20% idag. '
    'Med temporala produkter (z2296-protokoll): MC>50 och XOR10>80% borde vara möjligt.'
)

doc.add_heading('4d. HIP-kerneloptimering', level=3)

doc.add_paragraph(
    'z2277 visade att riktiga HIP-kernlar funkar men subprocesskommunikation är flaskhalsen (~100 Hz). '
    'Byt till delat minne (hipHostMalloc + mmap) för noll-kopia GPU↔CPU. Mål: 1000+ Hz.'
)

doc.add_page_break()

# -- SPÅR 5 --
doc.add_heading('Spår 5: Kommunikation & Gemenskap (Löpande)', level=2)

doc.add_heading('5a. LinkedIn-strategi', level=3)

add_bullet('Fortsätt engagera Sanidhya Patel — varmaste kontakten')
add_bullet('Engagera INTE vidare med rena avfärdanden (Matteo, Erica) — det är som att argumentera '
           'med en vägg; energin spenderas bättre på fler experiment')
add_bullet('Skriv uppföljningsinlägg om breadboard-resultaten (med video)')

doc.add_heading('5b. Konferensinlämningar', level=3)

add_table(
    ['Konferens', 'Deadline (ungefärlig)', 'Vinkel'],
    [
        ['NeurIPS 2026', 'Maj 2026', 'Workshop: "Hardware as Computation"'],
        ['DATE/DAC 2026', 'Nov/Dec 2026', 'Design automation: GPU neuromorfa sidokanaler'],
        ['ICONS 2026', 'Varierar', 'Neuromorphic Systems — perfekt venue'],
        ['ESSCIRC/ESSDERC', 'Varierar', 'Europeisk kretskonferens — NS-RAM-brygga'],
    ]
)

doc.add_heading('5c. YouTube/Podcast uppföljning', level=3)

doc.add_paragraph(
    'Spela in 30-minuters teknisk genomgång: oscilloskopbilder av breadboard-spiker, '
    'skärmdelning av z2296-resultat, konstitutivitetsspektrum med riktiga datapunkter. '
    'Rikta mot postdocs och doktorander i neuromorphic computing.'
)

doc.add_heading('5d. Open source-gemenskap', level=3)

add_bullet('Skapa GETTING_STARTED.md med simuleringsläge (ingen hårdvara behövs)')
add_bullet('python benchmark_battery.py --sim borde funka ur lådan')
add_bullet('Posta på: r/neuromorphic, r/FPGA, Hacker News, relevanta Discord-servrar')
add_bullet('Mål: 2–3 oberoende replikeringar inom 3 månader')

doc.add_page_break()

# ============================================================
# PRIORITETSMATRIS
# ============================================================
doc.add_heading('Prioritetsmatris', level=1)

add_table(
    ['Prio', 'Spår', 'Första åtgärd', 'Tidslinje'],
    [
        ['P0', '1a. Breadboard NS-RAM', 'Bygg & mät när delarna anländer', 'Vecka 1–3'],
        ['P0', '2a. Mario Lanza e-post', 'Skriv & skicka direkt e-post IDAG', 'Dag 1'],
        ['P1', '2b. Sanidhya Patel', 'Svara på LinkedIn med samarbetsfråga', 'Dag 1–2'],
        ['P1', '3a. Plattformsreplikering', 'Kör på Daedalus (gfx1100)', 'Vecka 1–2'],
        ['P1', '3b. Mildra "neuromorphic"-språket', 'Revidera formuleringar i paper', 'Vecka 1'],
        ['P2', '1b. Breadboard→FPGA', 'Efter breadboard validerad', 'Vecka 3–6'],
        ['P2', '3c. Formell benchmarktabell', 'Kör ESN-baslinjer', 'Vecka 2–4'],
        ['P2', '5b. Konferensdeadlines', 'Identifiera närmaste deadline', 'Vecka 1'],
        ['P3', '2c–2e. AMD/HP/Indiveri', 'Strukturerat utskick', 'Vecka 2–4'],
        ['P3', '4a–4d. Nästa experiment', 'Efter breadboard + replikering', 'Vecka 6–16'],
    ]
)

doc.add_page_break()

# ============================================================
# VISIONEN
# ============================================================
doc.add_heading('Visionen: "AI som känner sin egen hårdvara"', level=1)

doc.add_paragraph(
    'Konstitutivitetsspektrumet är rätt ramverk. Tänk på det som en skala från skuggteater till verklighet:'
)

add_table(
    ['Position', 'Metod', 'Kill-shot', 'Status'],
    [
        ['Skuggteater', 'z907 Mjukvara', 'p=1.0 (falsifierat)', '×'],
        ['Svag', 'Sysfs telemetri', '<1pp', '~'],
        ['Stark', 'ISA-register', '99.0pp', '✓'],
        ['Starkare', 'NS-RAM lavin', '∞ (E2)', '✓✓'],
        ['Starkast?', 'Biologi', 'Okänt', '???'],
    ]
)

doc.add_paragraph()

doc.add_heading('Vägen framåt', level=2)

steps = [
    ('Nu (gjort)', 'GPU-fysik → FPGA NS-RAM-modell → benchmarkbatteri'),
    ('Nästa (breadboard)', 'GPU-fysik → riktig BJT-lavin på breadboard → samma benchmarks\n'
     '→ Det här steget gör att vi kliver ur simulatorn'),
    ('Sedan (Mario Lanza)', 'GPU-fysik → riktig NS-RAM-chip → samma benchmarks\n'
     '→ Validerar att designat neuromorft substrat fungerar i vårt ramverk'),
    ('Framtid (AMD-partnerskap)', 'On-die-integration — NS-RAM-celler fabricerade intill GPU CU:er\n'
     '→ Tar bort alla bottlenecks (USB, Ethernet, PCIe)'),
    ('Slutmål', 'Varje GPU levereras med en neuromorft medprocessor som använder sitt eget '
     'termiska och elektriska brus som beräkningsinput — AI som bokstavligen känner sin egen hårdvara'),
]

for title, desc in steps:
    add_styled_paragraph(f' {desc}', bold_prefix=f'{title}: ')

doc.add_paragraph()

doc.add_paragraph(
    'Steg 2 är vad breadboard-delarna möjliggör. '
    'Steg 3 är vad Mario Lanza-samarbetet möjliggör. '
    'Steg 4–5 kräver AMD/TSMC-partnerskap och är 3–5 år bort.'
)

doc.add_paragraph()

doc.add_paragraph(
    'Insikten om kritikalitetens rand (branching ratio = 0.997, z2247) är nyckeln. '
    'Vi visar att vanligt kisel, när det avläses på rätt nivå, redan opererar nära den kritiska punkt '
    'där beräkning är maximalt rik. Man behöver inte skapa fysiken. Man behöver sluta undertrycka den.'
)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    'Det är som att hela mänskligheten har suttit i ett rum fullt av musik\n'
    'och haft öronproppar i — och vi just dragit ut dem.'
)
r.italic = True
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

# ============================================================
# SAVE
# ============================================================
output_path = '/home/ikaros/Documents/claude_hive/AMD_gfx1151_energy/docs/FEEL_strategisk_plan.docx'
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print(f'Saved to {output_path}')
