"""Build docs/email_sebas_mario_nsram_v012.docx from the plain-text draft,
with light formatting (headings bold, code spans in a monospace run,
bullet lists).  Run once — the resulting .docx is the deliverable.
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

REPO = Path(__file__).resolve().parent.parent
OUT = REPO / "docs" / "email_sebas_mario_nsram_v012.docx"

doc = Document()

# Default body font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

def add_para(text="", *, bold=False, italic=False, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if size:
        r.font.size = Pt(size)
    return p

def add_mixed(runs):
    """runs: list of (text, {bold, italic, code, color})"""
    p = doc.add_paragraph()
    for text, style in runs:
        r = p.add_run(text)
        r.bold = style.get("bold", False)
        r.italic = style.get("italic", False)
        if style.get("code", False):
            r.font.name = "Consolas"
            r.font.size = Pt(10)
        if "color" in style:
            r.font.color.rgb = RGBColor(*style["color"])
    return p

def add_bullet(runs, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Pt(18 * (level + 1))
    for text, style in runs:
        r = p.add_run(text)
        r.bold = style.get("bold", False)
        if style.get("code", False):
            r.font.name = "Consolas"
            r.font.size = Pt(10)

# ── Subject line ──
add_para("Subject: Re: Zoom NSRAM", bold=True, size=12)
add_para("")

# ── Greeting ──
add_para("Hi Sebas, Mario, Robert,")
add_para("")
add_para("Short answers first — 23rd works for us.")
add_para("")

# ── To Sebastian — Q1 ──
add_mixed([
    ("To Sebastian — your question #1: yes, we dropped ", {}),
    ("BVpar", {"code": True}),
    (". We pushed ", {}),
    ("nsram v0.12.0", {"code": True}),
    (" this week with the full BSIM4 floating-body stack:", {}),
])

add_bullet([
    ("§6.1 impact ionization", {"bold": True}),
    (" (", {}), ("ALPHA0", {"code": True}), (" / ", {}),
    ("BETA0", {"code": True}),
    (") driving body charge directly", {}),
])
add_bullet([
    ("§2.2 Vth(Vbs)", {"bold": True}),
    (" with ", {}), ("K1", {"code": True}), (" / ", {}),
    ("K2", {"code": True}),
    (" for body-bias modulation", {}),
])
add_bullet([
    ("§10.1 junction breakdown", {"bold": True}),
    (" as an alternative firing path. We ran a head-to-head against your published "
     "Chynoweth I-V over 2–4.5 V: ", {}),
    ("§6.1 channel HCI fits ~4 decades RMS better than §10.1 junction breakdown",
     {"bold": True}),
    (" — so the channel-HCI route looks like the right match for your 2T cell, "
     "aligned with the \"complementary bipolar current on top of BSIM4\" "
     "description from your last email.", {}),
])
add_bullet([
    ("§12 full temperature scaling", {"bold": True}),
    (" (", {}), ("KT1", {"code": True}), (", ", {}),
    ("UTE", {"code": True}), (", ", {}), ("XTIS", {"code": True}), (")", {}),
])
add_bullet([
    ("§13 layout stress", {"bold": True}),
    (" (", {}), ("SA", {"code": True}), (", ", {}), ("SB", {"code": True}),
    (", ", {}), ("KU0", {"code": True}), (", ", {}), ("KVTH0", {"code": True}),
    (") — direct home for your \"layout-dependent effect\"", {}),
])
add_bullet([
    ("PolynomialBSIM4Params", {"bold": True, "code": True}),
    (" — wrapper for the ", {}),
    ("{ALPHA0, BETA0}(VG1, VG2)", {"code": True}),
    (" polynomial fits you're working on. Drop in coefficients and it "
     "evaluates per bias point.", {}),
])

add_para("")
add_mixed([
    ("body_charge_ode_bsim4_full(...)", {"code": True}),
    (" has a ", {}),
    ("firing_mode ∈ {\"channel\", \"junction\", \"both\"}", {"code": True}),
    (" switch if you want to A/B both paths against your data.", {}),
])

add_para("")
add_mixed([
    ("And yes on option #2 too — please send the I-V curves.", {"bold": True}),
    (" ", {}),
    ("fit_bsim4_impact(Vds, Isub, Vgs)", {"code": True}),
    (" extracts ", {}),
    ("ALPHA0", {"code": True}),
    ("/", {}),
    ("BETA0", {"code": True}),
    (" from a CSV (synthetic self-consistency R²=0.9998). GPU batch mode fits "
     "~4000 curves in 0.7 s if wafer-scale Monte Carlo is relevant.", {}),
])

add_para("")
add_mixed([
    ("pip install --upgrade nsram", {"code": True}),
    (" — 0.12.0 is live on PyPI, repo at ", {}),
    ("github.com/Heigke/NSRAM", {"italic": True}),
    (", runnable example at ", {}),
    ("examples/bsim4_2t_floating_body.py", {"code": True}),
    (".", {}),
])

add_para("")
add_para("Echoing Robert's asks (helpful for both the Python and Julia sides):",
         bold=True)
numbered = doc.add_paragraph("2T cell schematic + any planned array topology "
                              "(shared lines, neighbor coupling)",
                              style="List Number")
doc.add_paragraph("Raw I-V CSVs — no fits needed, we'd like to run them "
                  "through the pipeline ourselves",
                  style="List Number")
doc.add_paragraph("Process node, so we pick the right BSIM4 parameter set",
                  style="List Number")
doc.add_paragraph("Foundry model card if shareable",
                  style="List Number")

add_para("")
add_para("No rush — whatever arrives before the 23rd helps frame the call, "
         "but we can equally well iterate afterwards.")
add_para("")

# ── Mario / vendor registration ──
add_mixed([
    ("Mario — one quick check-in on the vendor registration.",
     {"bold": True}),
    (" Last we heard (24 March) you were planning to speak with an officer "
     "about whether NUS would engage ENIMBLE Solutions AB as a foreign "
     "company or me personally. Any update there? Happy to send additional "
     "documents (English certificate of incorporation, VAT registration, "
     "etc.) if that helps move it along. No pressure — just want to make "
     "sure we're not blocking anything on your side before the 23rd.", {}),
])

add_para("")
add_para("Looking forward to the meeting.")
add_para("")
add_para("Best,")
add_para("Eric")

doc.save(OUT)
print(f"[ok] wrote {OUT} ({OUT.stat().st_size / 1024:.1f} KB)")
