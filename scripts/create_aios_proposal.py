#!/usr/bin/env python3
"""Generate AI-OS Project Proposal & NDA Word Document for ENIMBLE + HP.
Uses matplotlib for professional figures embedded as PNG.
FEEL is presented as ONE building block (HW interface), not the whole solution."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch
import numpy as np
import os
import tempfile

TMPDIR = tempfile.mkdtemp()
HP_BLUE = '#003D6B'
HP_LIGHT = '#0096D6'
GRAY = '#888888'
DARK = '#333333'

def fig_path(name):
    return os.path.join(TMPDIR, f'{name}.png')

# ═══════════════════════════════════════════════════════════════════════
# FIGURE GENERATION
# ═══════════════════════════════════════════════════════════════════════

def make_fig1_today_stack():
    """Figure 1: Today's fragmented computer stack."""
    fig, ax = plt.subplots(figsize=(7.5, 4.5))
    ax.set_xlim(0, 10); ax.set_ylim(0, 7); ax.axis('off')
    layers = [
        (0.5, 'HARDWARE', 'CPU, GPU, DRAM, sensors', '#D32F2F', 'Rich data — nobody reads it'),
        (1.5, 'FIRMWARE', 'PSP, SMU, BIOS', '#E64A19', 'Locked and opaque'),
        (2.5, 'DRIVERS', 'GPU, NIC, storage', '#F57C00', 'Fixed policies'),
        (3.5, 'OS / KERNEL', 'Linux, Windows', '#FFA000', 'Generic power governors'),
        (4.5, 'APPLICATIONS', 'Chrome, Excel, VS Code', '#7B1FA2', 'No idea about hardware'),
        (5.5, 'HUMAN', '"I want to save energy"', HP_BLUE, 'Intent lost in translation'),
    ]
    for y, label, detail, color, note in layers:
        box = FancyBboxPatch((0.8, y), 5.4, 0.8, boxstyle="round,pad=0.05",
                             facecolor=color, edgecolor='white', linewidth=1.5, alpha=0.9)
        ax.add_patch(box)
        ax.text(3.5, y+0.4, label, ha='center', va='center', fontsize=11, fontweight='bold', color='white')
        ax.text(6.5, y+0.55, detail, ha='left', va='center', fontsize=8, color=DARK)
        ax.text(6.5, y+0.25, note, ha='left', va='center', fontsize=7, color=GRAY, style='italic')
    for y in [1.3, 2.3, 3.3, 4.3, 5.3]:
        ax.plot([3.5, 3.5], [y, y+0.2], color='#CCCCCC', lw=1, ls='dotted')
        ax.text(0.35, y+0.08, 'X', ha='center', va='center', fontsize=8, color='#CC0000', fontweight='bold')
    ax.text(0.35, 3.5, 'NO\nINTELLIGENCE\nCONNECTS\nTHESE\nLAYERS', ha='center', va='center',
            fontsize=6.5, color='#CC0000', fontweight='bold',
            bbox=dict(boxstyle='round,pad=0.3', facecolor='#FFF0F0', edgecolor='#CC0000', alpha=0.8))
    fig.suptitle("Today's Computer: Every Layer Is Isolated", fontsize=13, fontweight='bold', color=HP_BLUE, y=0.97)
    plt.tight_layout(rect=[0,0,1,0.94])
    fig.savefig(fig_path('fig1'), dpi=200, bbox_inches='tight', facecolor='white'); plt.close(fig)


def make_fig2_aios_vision():
    """Figure 2: AI-OS unified — 5 building blocks."""
    fig, ax = plt.subplots(figsize=(8, 6))
    ax.set_xlim(0, 12); ax.set_ylim(0, 9.5); ax.axis('off')

    # Human on top
    box = FancyBboxPatch((2, 7.5), 8, 1.3, boxstyle="round,pad=0.1",
                         facecolor='#F5F5F5', edgecolor=HP_BLUE, linewidth=2)
    ax.add_patch(box)
    ax.text(6, 8.4, 'HUMAN', ha='center', va='center', fontsize=13, fontweight='bold', color=HP_BLUE)
    ax.text(6, 7.9, 'Natural language  •  Any task  •  Any complexity', ha='center', fontsize=8, color=GRAY)

    # AI-OS brain in the middle
    box = FancyBboxPatch((2, 5.2), 8, 1.7, boxstyle="round,pad=0.1",
                         facecolor=HP_BLUE, edgecolor=HP_LIGHT, linewidth=3, alpha=0.95)
    ax.add_patch(box)
    ax.text(6, 6.5, 'AI-OS REASONING LAYER', ha='center', fontsize=14, fontweight='bold', color='white')
    ax.text(6, 5.9, 'LLM  •  Task Planning  •  Tool Use  •  Context  •  Memory', ha='center', fontsize=8, color='#B0D4F1')
    ax.text(6, 5.55, 'Understands intent, plans actions, delegates to the right engine', ha='center', fontsize=7, color='#90C4E1')

    # Arrow human <-> reasoning
    ax.annotate('', xy=(6, 6.9), xytext=(6, 7.5), arrowprops=dict(arrowstyle='<->', color=HP_LIGHT, lw=3))

    # 5 building blocks at bottom
    blocks = [
        (0.3, 'HW\nINTERFACE', '(FEEL)', '#1B5E20', ['ISA registers', 'Sensor fusion', 'HW fingerprints']),
        (2.65, 'ENERGY\nMANAGER', '', '#B71C1C', ['DVFS control', 'Power budget', 'Thermal mgmt']),
        (5.0, 'APP\nORCHESTRATOR', '', '#4A148C', ['Launch apps', 'File/network', 'OS services']),
        (7.35, 'KNOWLEDGE\nBASE', '', '#01579B', ['RAG + memory', 'User prefs', 'System state']),
        (9.7, 'SECURITY\nLAYER', '', '#E65100', ['Permissions', 'Isolation', 'Audit trail']),
    ]
    for x, title, sub, color, items in blocks:
        box = FancyBboxPatch((x, 0.3), 2.0, 4.2, boxstyle="round,pad=0.08",
                             facecolor=color, edgecolor='white', linewidth=1.5, alpha=0.9)
        ax.add_patch(box)
        ax.text(x+1.0, 4.0, title, ha='center', va='center', fontsize=8.5, fontweight='bold', color='white')
        if sub:
            ax.text(x+1.0, 3.3, sub, ha='center', va='center', fontsize=7, color='#AADDAA')
        for i, item in enumerate(items):
            ax.text(x+1.0, 2.6 - i*0.55, item, ha='center', va='center', fontsize=7, color='#E0E0E0')
        # Arrow up to reasoning
        ax.annotate('', xy=(x+1.0, 5.2), xytext=(x+1.0, 4.5),
                    arrowprops=dict(arrowstyle='->', color=HP_LIGHT, lw=1.8))

    fig.suptitle('AI-OS: Five Building Blocks Under One Brain', fontsize=13, fontweight='bold', color=HP_BLUE, y=0.98)
    plt.tight_layout(rect=[0,0,1,0.95])
    fig.savefig(fig_path('fig2'), dpi=200, bbox_inches='tight', facecolor='white'); plt.close(fig)


def make_fig3_flow():
    """Figure 3: How a request flows through AI-OS."""
    fig, ax = plt.subplots(figsize=(8, 3.5))
    ax.set_xlim(0, 12); ax.set_ylim(0, 4); ax.axis('off')

    steps = [
        (0.3, 'User\nRequest', '#F5F5F5', HP_BLUE),
        (2.4, 'Reasoning\nLayer', HP_BLUE, 'white'),
        (4.5, 'Delegate to\nEngines', '#4A148C', 'white'),
        (6.6, 'HW Sense\n+ Actuate', '#1B5E20', 'white'),
        (8.7, 'Execute\n+ Optimize', '#B71C1C', 'white'),
        (10.3, 'Report\nBack', '#F5F5F5', HP_BLUE),
    ]
    for x, label, bg, fg in steps:
        w = 1.7 if x < 10 else 1.4
        box = FancyBboxPatch((x, 1.0), w, 2.0, boxstyle="round,pad=0.1",
                             facecolor=bg, edgecolor=HP_BLUE if bg=='#F5F5F5' else 'white', linewidth=1.5)
        ax.add_patch(box)
        ax.text(x+w/2, 2.0, label, ha='center', va='center', fontsize=8.5, fontweight='bold', color=fg)
    # Arrows
    for x in [2.0, 4.1, 6.2, 8.3, 10.0]:
        ax.annotate('', xy=(x+0.3, 2.0), xytext=(x, 2.0),
                    arrowprops=dict(arrowstyle='->', color=HP_LIGHT, lw=2))

    ax.text(1.15, 0.5, '"Investigate DRAM\nneuromorphic comp."', ha='center', fontsize=7, color=GRAY, style='italic')
    ax.text(3.25, 0.5, 'Break into\nsub-tasks', ha='center', fontsize=7, color=GRAY)
    ax.text(5.35, 0.5, 'HW interface,\nenergy, apps', ha='center', fontsize=7, color=GRAY)
    ax.text(7.45, 0.5, 'Read sensors,\noptimize power', ha='center', fontsize=7, color=GRAY)
    ax.text(9.55, 0.5, 'Run sims,\nwrite report', ha='center', fontsize=7, color=GRAY)
    ax.text(10.95, 0.5, 'Plain language\nsummary', ha='center', fontsize=7, color=GRAY)

    fig.suptitle('How a Request Flows Through AI-OS', fontsize=12, fontweight='bold', color=HP_BLUE, y=0.98)
    plt.tight_layout(rect=[0,0,1,0.92])
    fig.savefig(fig_path('fig3'), dpi=200, bbox_inches='tight', facecolor='white'); plt.close(fig)


def make_fig4_feel_summary():
    """Figure 4: FEEL as the HW interface — key stats."""
    fig, axes = plt.subplots(1, 3, figsize=(8, 3.2))

    # Pass rates
    exps = ['z2060', 'z2068', 'z2076', 'z2090', 'z2095']
    rates = [100, 95.5, 100, 94.4, 95.0]
    colors = [HP_LIGHT]*3 + [HP_BLUE, HP_BLUE]
    axes[0].barh(exps, rates, color=colors, edgecolor='white', height=0.6)
    axes[0].set_xlim(0, 112)
    axes[0].set_xlabel('Pass Rate (%)', fontsize=8)
    axes[0].set_title('Test Pass Rates', fontsize=10, fontweight='bold', color=HP_BLUE)
    for i, v in enumerate(rates):
        axes[0].text(v+1, i, f'{v:.0f}%', va='center', fontsize=7)
    axes[0].tick_params(labelsize=7)

    # Kill-shot
    labels = ['Vision\nAccuracy', 'Language\nPerplexity']
    values = [99.0, 28.3]
    cols = ['#D32F2F', '#E64A19']
    bars = axes[1].bar(labels, values, color=cols, edgecolor='white', width=0.5)
    axes[1].set_ylabel('% Degradation', fontsize=8)
    axes[1].set_title('Kill-Shot: HW is Causal', fontsize=10, fontweight='bold', color=HP_BLUE)
    for bar, v in zip(bars, values):
        axes[1].text(bar.get_x()+bar.get_width()/2, v+1.5, f'{v}%', ha='center', fontsize=9, fontweight='bold')
    axes[1].tick_params(labelsize=7)

    # Energy
    cats = ['Fixed\nDVFS', 'AI Self-\nOptimized']
    vals = [100, 87.2]
    cols = ['#CCCCCC', '#1B5E20']
    bars = axes[2].bar(cats, vals, color=cols, edgecolor='white', width=0.5)
    axes[2].set_ylabel('Relative Energy (%)', fontsize=8)
    axes[2].set_title('Energy Savings', fontsize=10, fontweight='bold', color=HP_BLUE)
    axes[2].set_ylim(0, 115)
    axes[2].axhline(y=100, color='#CCCCCC', ls='--', lw=0.5)
    for bar, v in zip(bars, vals):
        axes[2].text(bar.get_x()+bar.get_width()/2, v+1.5,
                     f'{v}%' if v<100 else 'baseline', ha='center', fontsize=8,
                     fontweight='bold' if v<100 else 'normal')
    axes[2].tick_params(labelsize=7)

    fig.suptitle('FEEL: Proven Hardware Interface Technology', fontsize=12, fontweight='bold', color=HP_BLUE, y=1.02)
    plt.tight_layout()
    fig.savefig(fig_path('fig4'), dpi=200, bbox_inches='tight', facecolor='white'); plt.close(fig)


def make_fig5_timeline():
    """Figure 5: Gantt-style timeline."""
    fig, ax = plt.subplots(figsize=(8, 3))
    ax.set_xlim(0, 24); ax.set_ylim(-0.5, 5.5); ax.axis('off')
    phases = [
        (0, 'Phase 0: Alignment & NDA', 0, 2, '#9E9E9E'),
        (1, 'Phase 1: HW Interface (FEEL port)', 2, 6, '#1B5E20'),
        (2, 'Phase 2: Energy + Knowledge', 5, 9, '#B71C1C'),
        (3, 'Phase 3: Reasoning + UI', 8, 13, '#4A148C'),
        (4, 'Phase 4: Integration & Testing', 13, 17, HP_BLUE),
        (5, 'Phase 5: Productization', 16, 22, HP_LIGHT),
    ]
    for row, label, start, end, color in phases:
        y = 5 - row
        ax.barh(y, end-start, left=start, height=0.6, color=color, edgecolor='white', lw=1.5, alpha=0.9)
        ax.text(start+(end-start)/2, y, label, ha='center', va='center', fontsize=7.5, fontweight='bold', color='white')
    months = ['Mar','Jun','Sep','Dec','Mar','Jun','Sep','Dec']
    for i, m in enumerate(months):
        ax.text(i*3, -0.4, m, ha='center', fontsize=7, color=GRAY)
    ax.text(0, -0.7, '2026', ha='left', fontsize=8, fontweight='bold', color=DARK)
    ax.text(12, -0.7, '2027', ha='left', fontsize=8, fontweight='bold', color=DARK)
    ax.axvline(x=12, color='#EEEEEE', ls='-', lw=0.5)
    ax.set_title('Project Timeline (Approximate)', fontsize=11, fontweight='bold', color=HP_BLUE, pad=10)
    plt.tight_layout()
    fig.savefig(fig_path('fig5'), dpi=200, bbox_inches='tight', facecolor='white'); plt.close(fig)


def make_fig6_building_blocks():
    """Figure 6: What each building block does (simple grid)."""
    fig, ax = plt.subplots(figsize=(8, 4))
    ax.set_xlim(0, 10); ax.set_ylim(0, 6); ax.axis('off')

    blocks = [
        (0.2, 3.2, 'HW Interface\n(FEEL)', '#1B5E20',
         'Talks to GPU registers,\nsensors, firmware.\nReads 900+ values/sec.\nProven in 38+ experiments.'),
        (2.2, 3.2, 'Energy Manager', '#B71C1C',
         'Controls DVFS, power\nbudgets, thermal limits.\n12.8% savings proven.\nAdapts per workload.'),
        (4.2, 3.2, 'App Orchestrator', '#4A148C',
         'Launches apps, manages\nfiles, network, storage.\nBridges OS services\nto the AI reasoning.'),
        (6.2, 3.2, 'Knowledge Base', '#01579B',
         'RAG, user preferences,\nsystem state history.\nLearns from usage.\nPersonalizes over time.'),
        (8.2, 3.2, 'Security Layer', '#E65100',
         'Permissions, sandboxing,\naudit trail, isolation.\nUser stays in control.\nNo unauthorized actions.'),
    ]
    for x, y, title, color, desc in blocks:
        box = FancyBboxPatch((x, y-2.8), 1.8, 2.6, boxstyle="round,pad=0.08",
                             facecolor=color, edgecolor='white', linewidth=1.5, alpha=0.9)
        ax.add_patch(box)
        ax.text(x+0.9, y, title, ha='center', va='center', fontsize=8, fontweight='bold', color='white')
        ax.text(x+0.9, y-1.4, desc, ha='center', va='center', fontsize=6.5, color='#E0E0E0')

    # "FEEL is HERE" callout on first block
    ax.annotate('FEEL provides\nthis piece', xy=(1.1, 0.4), xytext=(1.1, -0.2),
                fontsize=8, fontweight='bold', color='#1B5E20', ha='center',
                arrowprops=dict(arrowstyle='->', color='#1B5E20', lw=1.5))

    ax.text(5.0, 5.5, 'Five Building Blocks — Each Solves One Problem', ha='center',
            fontsize=12, fontweight='bold', color=HP_BLUE)
    plt.tight_layout()
    fig.savefig(fig_path('fig6'), dpi=200, bbox_inches='tight', facecolor='white'); plt.close(fig)


# ═══════════════════════════════════════════════════════════════════════
# Generate all figures
# ═══════════════════════════════════════════════════════════════════════
print("Generating figures...")
make_fig1_today_stack()
make_fig2_aios_vision()
make_fig3_flow()
make_fig4_feel_summary()
make_fig5_timeline()
make_fig6_building_blocks()
print("Figures done.")

# ═══════════════════════════════════════════════════════════════════════
# DOCUMENT HELPERS
# ═══════════════════════════════════════════════════════════════════════
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x00, 0x3D, 0x6B)
    hs.font.name = 'Calibri'

def add_hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '12'); b.set(qn('w:space'), '1'); b.set(qn('w:color'), '003D6B')
    pBdr.append(b); pPr.append(pBdr)

def shade(cell, color):
    s = OxmlElement('w:shd'); s.set(qn('w:fill'), color); s.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(s)

def stbl(doc, headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            for r in p.runs: r.bold = True; r.font.color.rgb = RGBColor(255,255,255); r.font.size = Pt(10)
        shade(c, '003D6B')
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = str(val)
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(10)
            if ri % 2 == 0: shade(c, 'E8F0FE')
    tblPr = t._tbl.tblPr if t._tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for e in ('top','left','bottom','right','insideH','insideV'):
        el = OxmlElement(f'w:{e}'); el.set(qn('w:val'),'single'); el.set(qn('w:sz'),'4'); el.set(qn('w:color'),'999999')
        borders.append(el)
    tblPr.append(borders)
    return t

def afig(doc, name, caption, width=6.0):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(fig_path(name), width=Inches(width))
    c = doc.add_paragraph(caption); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in c.runs: r.italic = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x66,0x66,0x66)

def bold_para(doc, title, text):
    p = doc.add_paragraph(); r = p.add_run(title + ': '); r.bold = True; p.add_run(text)

# ═══════════════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════════════
for _ in range(4): doc.add_paragraph()
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('AI-OS'); r.font.size = Pt(42); r.bold = True; r.font.color.rgb = RGBColor(0x00,0x3D,0x6B)
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run('The AI Operating System'); r.font.size = Pt(22); r.font.color.rgb = RGBColor(0x00,0x96,0xD6)
doc.add_paragraph()
tg = doc.add_paragraph(); tg.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tg.add_run('A unified intelligence layer from silicon to human intent')
r.font.size = Pt(14); r.italic = True; r.font.color.rgb = RGBColor(0x55,0x55,0x55)
for _ in range(4): doc.add_paragraph()
info = doc.add_paragraph(); info.alignment = WD_ALIGN_PARAGRAPH.CENTER
for line in ['Project Proposal & Collaboration Framework', '',
             'ENIMBLE AB', "(Eric's Neuromapping and Impulse response",
             'for Better Life Enhancement Solutions AB)', 'Eric Bergvall, Founder', '',
             'HP Inc.', 'Albert Bruhner, Contact', '', 'February 2026', '', 'CONFIDENTIAL']:
    r = info.add_run(line + '\n'); r.font.size = Pt(12)
    if line in ('ENIMBLE AB', 'HP Inc.', 'CONFIDENTIAL'): r.bold = True
    if line == 'CONFIDENTIAL': r.font.color.rgb = RGBColor(0xCC,0x00,0x00)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# TOC
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
for item in ['1. Executive Summary', '2. The Problem', '3. The Vision: AI-OS',
             '4. The Five Building Blocks', '5. FEEL: The Hardware Interface (Proven Technology)',
             '6. How It All Works Together', '7. Use Cases',
             '8. Project Plan', '9. What ENIMBLE Needs from HP',
             '10. Draft NDA Terms', '11. Team & Next Steps']:
    p = doc.add_paragraph(item); p.paragraph_format.space_after = Pt(2)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('1. Executive Summary', level=1); add_hr(doc)

doc.add_paragraph(
    'AI-OS is a project to build an AI operating system \u2014 a single intelligence layer '
    'that connects the human user to the hardware, through every abstraction level in between. '
    'It is not one technology; it is five building blocks working together:'
)
for b in [
    'A hardware interface that reads and controls GPU/CPU state in real-time',
    'An energy manager that self-optimizes power, thermals, and performance',
    'An app orchestrator that launches tools and manages system services',
    'A knowledge base that remembers user preferences and system history',
    'A security layer that keeps the user in control at all times',
]:
    doc.add_paragraph(b, style='List Bullet')

doc.add_paragraph(
    'ENIMBLE has already built and proven the first block: the FEEL project demonstrated '
    'in 38+ experiments that neural networks can read and write GPU hardware registers during '
    'their own computation, achieving 12.8% energy savings and 24.5% language model improvement. '
    'This gives AI-OS its "body" \u2014 the ability to sense and control the physical machine.'
)
doc.add_paragraph(
    'HP brings the hardware platforms, OEM integration, and route to market. '
    'Together, we build the remaining four blocks around the proven FEEL foundation, '
    'creating a computer that understands whatever the person wants \u2014 from deep '
    'technical research to simple tax spreadsheets \u2014 and handles it intelligently.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 2. THE PROBLEM
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('2. The Problem', level=1); add_hr(doc)

doc.add_paragraph(
    'Modern computers have a fundamental disconnect. Each layer of the stack speaks '
    'a different language and has no idea what the other layers are doing.'
)
afig(doc, 'fig1', 'Figure 1: Today\'s computer \u2014 six layers, zero communication between them.')

doc.add_paragraph(
    'The GPU has hundreds of sensors (thermal, power, frequency, cache stats) '
    'but applications can\'t see them. The user says "save battery" but the hardware '
    'doesn\'t hear it. The OS uses generic power governors that know nothing about '
    'the actual workload. Every layer is siloed.'
)
doc.add_paragraph(
    'The result: wasted energy, suboptimal performance, and users who manually manage '
    'settings they shouldn\'t need to think about. No current AI assistant (Copilot, Siri, '
    'Google Assistant) reaches below the application layer \u2014 they\'re chatbots on top '
    'of the same disconnected stack.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 3. THE VISION
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('3. The Vision: AI-OS', level=1); add_hr(doc)

doc.add_paragraph(
    'AI-OS places a reasoning layer in the middle that talks to everything: '
    'hardware below, human above, and five specialized engines that handle '
    'the actual work.'
)
afig(doc, 'fig2', 'Figure 2: AI-OS architecture \u2014 five building blocks under one reasoning layer.', width=6.5)

doc.add_heading('What Makes This Different', level=2)
stbl(doc, ['Feature', 'Existing AI Assistants', 'AI-OS'],
    [['Hardware access', 'None \u2014 app layer only', 'Reads 900+ HW sensors, controls DVFS'],
     ['Energy awareness', 'None', 'Self-optimizes per workload (12.8% savings proven)'],
     ['Depth of control', 'Can open apps', 'Controls HW registers, OS, apps, and more'],
     ['Task complexity', 'Simple Q&A, search', '"Research DRAM neuromorphic" \u2192 full pipeline'],
     ['Self-awareness', 'None', 'Knows own thermal state, power draw, utilization'],
     ['Adaptation', 'Static', 'Learns from usage, optimizes over time'],
])
doc.add_paragraph()
doc.add_paragraph(
    'The key: AI-OS is not a chatbot bolted on top. It reaches into the hardware, '
    'understands the full stack, and adapts in real-time. FEEL gives it the "body"; '
    'the other four blocks give it the "brain", "tools", "memory", and "safety".'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 4. FIVE BUILDING BLOCKS
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('4. The Five Building Blocks', level=1); add_hr(doc)

doc.add_paragraph(
    'AI-OS follows a KISS principle: five simple blocks, each solving one problem. '
    'They communicate through the central reasoning layer.'
)
afig(doc, 'fig6', 'Figure 3: Five building blocks \u2014 FEEL provides the hardware interface (green, left). The other four need to be built.', width=6.5)

stbl(doc, ['Block', 'What It Does', 'Status', 'Key Challenge'],
    [['1. HW Interface (FEEL)', 'Reads/writes GPU registers, fuses sensor data', 'Proven (38+ experiments)', 'Port to HP platforms, extend to Intel'],
     ['2. Energy Manager', 'Controls DVFS, power budget, thermals', 'Partially proven (FEEL)', 'Cross-platform power APIs, workload prediction'],
     ['3. App Orchestrator', 'Launches apps, manages files/network/OS', 'To be built', 'OS integration, sandboxing, cross-platform'],
     ['4. Knowledge Base', 'RAG, user prefs, system state, learning', 'To be built', 'Privacy, efficient retrieval, personalization'],
     ['5. Security Layer', 'Permissions, isolation, audit, user control', 'To be built', 'Trust model, privilege escalation prevention'],
])

doc.add_paragraph()
doc.add_paragraph(
    'FEEL solves the hardest part \u2014 making AI talk to hardware. '
    'The remaining blocks use well-understood technologies (LLMs, RAG, '
    'OS APIs, sandboxing) applied in a new way. The innovation is in '
    'the integration, not in reinventing each piece.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 5. FEEL: THE HARDWARE INTERFACE
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('5. FEEL: The Hardware Interface (Proven Technology)', level=1); add_hr(doc)

doc.add_paragraph(
    'FEEL (Functionally Embodied Emergent Learning) is the proven technology that gives '
    'AI-OS its ability to sense and control hardware. Over 38+ experiments on AMD RDNA 3.5, '
    'ENIMBLE demonstrated that neural networks can:'
)
for b in [
    'Read and write ISA registers during their own forward pass',
    'Sense thermal, power, frequency, and memory fabric state in real-time',
    'Self-optimize energy consumption (12.8% better than any fixed setting)',
    'Improve language model quality through hardware awareness (24.5% perplexity improvement)',
]:
    doc.add_paragraph(b, style='List Bullet')

afig(doc, 'fig4', 'Figure 4: FEEL key results \u2014 high test pass rates, causal hardware dependence, and energy savings.', width=6.5)

doc.add_paragraph(
    'The "kill-shot" test is the most important result: scrambling the hardware state '
    'causes a 99.0 percentage-point accuracy drop in the vision model. This proves the '
    'AI genuinely depends on hardware state \u2014 it\'s not decorative.'
)
doc.add_paragraph(
    'In AI-OS, FEEL serves as the interface between the reasoning layer and the '
    'physical hardware. The LLM doesn\'t read registers directly \u2014 FEEL translates '
    'hardware state into a form the AI can understand and act on. Think of it as '
    'giving the computer a nervous system.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 6. HOW IT WORKS TOGETHER
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('6. How It All Works Together', level=1); add_hr(doc)

afig(doc, 'fig3', 'Figure 5: Request flow \u2014 from user intent to hardware actuation and back.', width=6.5)

doc.add_heading('Example A: Deep Technical Research', level=2)
p = doc.add_paragraph(); r = p.add_run('"I want to investigate how DRAM can be used for neuromorphic computation"'); r.italic = True
for t, d in [
    ('Reasoning Layer', 'Understands this is a multi-step research task. Plans: literature search, simulation setup, experiments, report.'),
    ('App Orchestrator', 'Opens browser, searches papers, downloads PDFs, sets up simulation environment.'),
    ('HW Interface (FEEL)', 'Reads GPU/CPU state: "32GB DRAM available, GPU at 45\u00b0C, 85W headroom."'),
    ('Energy Manager', 'Heavy compute coming \u2192 boosts clocks, monitors thermals, ensures stable power.'),
    ('Knowledge Base', 'Stores findings, remembers user\'s previous work on related topics.'),
    ('Security Layer', 'Verifies the simulations stay sandboxed, no unauthorized network access.'),
    ('Back to User', '"I found 12 relevant papers and ran 3 simulations. Here\'s what DRAM neuromorphic approaches look like..."'),
]:
    bold_para(doc, t, d)

doc.add_paragraph()
doc.add_heading('Example B: Simple Productivity', level=2)
p = doc.add_paragraph(); r = p.add_run('"I need an Excel sheet for my business tax reporting"'); r.italic = True
for t, d in [
    ('Reasoning Layer', 'Simple productivity task. Plans: create spreadsheet with categories and formulas.'),
    ('Energy Manager', 'Light workload \u2192 drops GPU to minimum, saves battery.'),
    ('App Orchestrator', 'Creates Excel file with Swedish tax categories, VAT formulas, formatting.'),
    ('Knowledge Base', 'Remembers user\'s company name and previous tax structure.'),
    ('Back to User', '"Here\'s your tax spreadsheet. I used the same format as last quarter."'),
]:
    bold_para(doc, t, d)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 7. USE CASES
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('7. Use Cases', level=1); add_hr(doc)

stbl(doc, ['Scenario', 'User Says', 'AI-OS Does'],
    [['Battery life', '"Save battery, I\'m on a flight"', 'HW Interface reads power state, Energy Manager drops clocks + dims screen'],
     ['Performance', '"Max performance for this render"', 'Energy Manager boosts clocks, HW Interface monitors thermals'],
     ['Research', '"Research quantum error correction"', 'Reasoning plans, App Orchestrator runs tools, Knowledge Base stores findings'],
     ['Productivity', '"Prepare my quarterly report"', 'App Orchestrator gathers data, creates charts, formats document'],
     ['Diagnostics', '"Why is my fan so loud?"', 'HW Interface reads thermals, Reasoning explains in plain language'],
     ['Development', '"Profile and optimize this code"', 'HW Interface reads perf counters, Reasoning identifies bottlenecks'],
     ['Maintenance', '"Is my SSD healthy?"', 'HW Interface reads SMART data, Knowledge Base tracks trends, warns if degrading'],
     ['Learning', '"Teach me how GPUs work"', 'Reasoning + Knowledge Base create personalized lesson, HW Interface shows live examples'],
])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 8. PROJECT PLAN
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('8. Project Plan', level=1); add_hr(doc)

afig(doc, 'fig5', 'Figure 6: Approximate project timeline.', width=6.5)

stbl(doc, ['Phase', 'Duration', 'What Gets Built', 'Key Activities'],
    [['Phase 0: Alignment', '1\u20132 mo', 'NDA, scope, target HW', 'Define IP boundaries, access needs, pick HP platforms'],
     ['Phase 1: HW Interface', '3\u20134 mo', 'FEEL ported to HP', 'Port 7-layer sensing, validate on HP AMD + Intel hardware'],
     ['Phase 2: Energy + Knowledge', '3\u20134 mo', 'Energy Manager + Knowledge Base', 'Cross-platform DVFS, RAG system, user preference engine'],
     ['Phase 3: Reasoning + UI', '3\u20134 mo', 'LLM integration + natural UI', 'Task planning, tool use, conversational interface'],
     ['Phase 4: Integration', '2\u20133 mo', 'All 5 blocks working together', 'End-to-end testing, demo scenarios, security hardening'],
     ['Phase 5: Product', '3\u20136 mo', 'Shippable AI-OS', 'OEM integration, user testing, documentation'],
])
doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('Total: approximately 15\u201323 months from NDA to shippable prototype.'); r.bold = True

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 9. WHAT ENIMBLE NEEDS
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('9. What ENIMBLE Needs from HP', level=1); add_hr(doc)

doc.add_heading('Hardware', level=2)
stbl(doc, ['Item', 'Why', 'Priority'],
    [['HP laptop (AMD APU)', 'Primary dev platform for HW Interface porting', 'Critical'],
     ['HP laptop (Intel)', 'Cross-platform validation', 'High'],
     ['HP server (optional)', 'Datacenter variant testing', 'Medium'],
])

doc.add_paragraph()
doc.add_heading('Technical Documentation', level=2)
stbl(doc, ['Item', 'Why', 'Priority'],
    [['Sensor documentation', 'Map 900+ sensors on HP platforms', 'Critical'],
     ['Power management APIs', 'Integrate Energy Manager with HP power framework', 'High'],
     ['BIOS/firmware docs', 'Unlock perf counters, fine DVFS control', 'High'],
     ['WMI/ACPI interfaces', 'Control HP-specific features (fans, power limits)', 'Medium'],
     ['Driver docs or source', 'Extend HW Interface to HP-specific hardware', 'Medium'],
])

doc.add_paragraph()
doc.add_heading('Collaboration', level=2)
for b in [
    'Bi-weekly technical meetings with HP firmware/driver engineers',
    'HP contact person (Albert Bruhner) for coordination',
    'Co-development agreement for joint IP on AI-OS innovations',
]:
    doc.add_paragraph(b, style='List Bullet')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 10. DRAFT NDA
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('10. Draft Non-Disclosure Agreement Terms', level=1); add_hr(doc)

p = doc.add_paragraph()
r = p.add_run('NOTE: Draft outline for discussion. Formal NDA to be prepared by legal counsel.')
r.italic = True; r.font.color.rgb = RGBColor(0xCC,0x00,0x00)

doc.add_heading('Parties', level=2)
doc.add_paragraph(
    'ENIMBLE AB (Eric\'s Neuromapping and Impulse response for Better Life Enhancement '
    'Solutions AB), represented by Eric Bergvall (Founder).')
doc.add_paragraph('HP Inc., represented by Albert Bruhner (Contact Person).')

doc.add_heading('Purpose', level=2)
doc.add_paragraph(
    'To evaluate and develop a joint project ("AI-OS") combining ENIMBLE\'s hardware-embodied '
    'AI technology with HP\'s hardware platforms, requiring exchange of confidential information.')

doc.add_heading('Confidential Information', level=2)
stbl(doc, ['From ENIMBLE', 'From HP'],
    [['FEEL research data, code, results', 'Hardware schematics and sensor topology'],
     ['AI-OS architecture designs', 'BIOS/firmware documentation'],
     ['Hardware sensing algorithms', 'Power management APIs'],
     ['Model weights and training methods', 'Product roadmap and specifications'],
     ['Experimental results (38+ experiments)', 'Driver source and internal tools'],
])

doc.add_paragraph()
doc.add_heading('Key Terms', level=2)
for t, d in [
    ('Duration', '3 years from signing, option to extend.'),
    ('Scope', 'AI-OS project only. No use for competing products.'),
    ('IP Ownership', 'Each party keeps pre-existing IP. Joint IP for co-developed work defined separately.'),
    ('Return of Materials', 'Return or destroy within 30 days of termination.'),
    ('Exceptions', 'Public information, independently developed, third-party received.'),
    ('ENIMBLE Access', 'Physical HP hardware, remote documentation access, regular engineer meetings. All access logged.'),
    ('Publication', 'ENIMBLE may publish FEEL methodology (no HP details) with 30-day HP review.'),
    ('Governing Law', 'To be determined (Swedish or US jurisdiction).'),
]:
    bold_para(doc, t, d)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 11. TEAM & NEXT STEPS
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('11. Team & Next Steps', level=1); add_hr(doc)

stbl(doc, ['Person', 'Organization', 'Role'],
    [['Eric Bergvall', 'ENIMBLE AB (Founder)', 'Project Lead, AI/Hardware Architecture'],
     ['Albert Bruhner', 'HP Inc.', 'HP Contact, Hardware Integration'],
])

doc.add_paragraph()
doc.add_heading('Eric Bergvall \u2014 Project Lead', level=2)
doc.add_paragraph(
    'Eric Bergvall is an AI architect whose career spans both ends of the intelligence spectrum: '
    'from hardware-level AI integration to deploying AI into human organizations and business processes. '
    'This dual expertise is exactly what AI-OS requires \u2014 bridging the gap from silicon to human intent.'
)
doc.add_heading('Key Background', level=3)
for b in [
    'MSc Engineering Physics \u2014 Machine Intelligence, Lund University (2014\u20132020)',
    'Systems Architect & Program Manager at Saab (2021\u20132025): Led the Central AI Initiative, '
    'architected the Defence Cloud, and led the 500+ person R&T AI & Autonomy Cluster. '
    'Built proofs-of-concept with Nvidia and Helsing. Designed HPC infrastructure for defence AI.',
    'AI Architect at The AI Framework / Advania (2025\u2013present): Helps businesses deploy AI '
    'on-premises, navigating proprietary data and trust requirements \u2014 the same trust challenges '
    'AI-OS must solve for end users.',
    'Founder of ENIMBLE AB: Built the FEEL project from scratch \u2014 38+ experiments, 5 architectures, '
    'world\'s first NN writing ISA registers during its own forward pass.',
    'Research background in Neural ODEs (Sentian.AI), Reinforcement Learning for industrial processes '
    '(Master\'s thesis), and bioinformatics neural networks (Biomedical Centre, Lund \u2014 Nature Science review).',
]:
    doc.add_paragraph(b, style='List Bullet')

doc.add_paragraph(
    'Eric\'s vision for AI-OS bridges three domains he has worked in throughout his career: '
    'the existing system domain (hardware, firmware, OS), the AI assistance domain (reasoning, '
    'optimization, adaptation), and the human domain (trust, usability, business value). '
    'Defence-grade systems thinking from Saab combined with real-world AI deployment experience '
    'at Advania gives AI-OS a uniquely grounded perspective.'
)

doc.add_paragraph()
doc.add_heading('ENIMBLE Track Record', level=2)
for b in [
    '38+ experiments across 5 neural architectures (CNN, ResNet, ViT, Transformer, GPT-2)',
    'World\'s first neural network writing ISA registers during its own forward pass',
    'World\'s first hardware-embodied language model (GPT-2 + regime-bound LoRA)',
    '12.8% energy savings through AI self-optimization (proven)',
    '24.5% language model perplexity improvement through hardware embodiment',
    '99.0 percentage-point kill-shot: scrambled HW destroys model \u2192 hardware is causal, not decorative',
    'All on commodity AMD hardware \u2014 no custom silicon required',
]:
    doc.add_paragraph(b, style='List Bullet')

doc.add_paragraph()
doc.add_heading('Next Steps', level=2)
stbl(doc, ['Step', 'Action', 'Owner', 'Target'],
    [['1', 'Review this proposal', 'HP (Albert Bruhner)', 'March 2026'],
     ['2', 'Technical deep-dive & FEEL demo', 'ENIMBLE + HP', 'March 2026'],
     ['3', 'NDA negotiation & signing', 'Legal (both)', 'April 2026'],
     ['4', 'HP hardware provisioning', 'HP', 'April 2026'],
     ['5', 'Phase 1 kickoff', 'ENIMBLE', 'May 2026'],
])

doc.add_paragraph(); doc.add_paragraph()
doc.add_heading('Contact', level=2)
p = doc.add_paragraph(); r = p.add_run('ENIMBLE AB\n'); r.bold = True
p.add_run('Eric Bergvall, Founder\nEmail: [to be added]\nPhone: [to be added]\n')
doc.add_paragraph()
p = doc.add_paragraph(); r = p.add_run('HP Inc.\n'); r.bold = True
p.add_run('Albert Bruhner\nEmail: [to be added]\n')
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('This document is confidential and intended solely for the named parties. '
              'Do not distribute without written permission from both ENIMBLE AB and HP Inc.')
r.italic = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x99,0x99,0x99)

# ═══════════════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════════════
out = '/home/ikaros/Documents/claude_hive/AMD_gfx1151_energy/results/AIOS_Project_Proposal_ENIMBLE_HP.docx'
doc.save(out)
print(f'Saved: {out}')
print(f'Size: {os.path.getsize(out)/1024:.1f} KB')
