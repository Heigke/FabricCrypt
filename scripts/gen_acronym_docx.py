#!/usr/bin/env python3
"""Generate FEEL acronym glossary as .docx"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

title = doc.add_heading('FEEL — Akronymer och Förklaringar', level=1)
doc.add_paragraph('Ordlista för FEEL-projektet (Functionally Embodied Emergent Learning)')

rows = [
    ("FEEL", "Functionally Embodied Emergent Learning", "Vårt ramverk: lärande som uppstår ur fysisk hårdvara, inte bara mjukvara"),
    ("NS-RAM", "Neuristive RAM (Neuro-Resistive RAM)", "Marios minnesceller som beter sig som neuroner — HfO₂-baserade memristorer"),
    ("LIF", "Leaky Integrate-and-Fire", "Enklaste neuronmodellen: samla ström, läck lite, skjut spike vid tröskel"),
    ("FPGA", "Field-Programmable Gate Array", "Omprogrammerbar krets — vi kör 128 LIF-neuroner på en Arty A7"),
    ("GPU", "Graphics Processing Unit", "AMD Radeon 8060S — vi använder den UNDER normalt lager"),
    ("HIP", "Heterogeneous-compute Interface for Portability", "AMD:s programmeringsspråk för GPU-beräkning (som CUDA för Nvidia)"),
    ("RDNA", "Radeon DNA", "AMD:s GPU-arkitekturfamilj — vår är RDNA4 (gfx1151)"),
    ("VRM", "Voltage Regulator Module", "Spänningsregulator på GPU:n — har naturligt 1/f-brus"),
    ("SMN", "System Management Network", "Intern buss i AMD-chip för firmware/temperatur/klockor"),
    ("SMU", "System Management Unit", "Processorn inuti GPU:n som styr spänning, temp, klockor"),
    ("PSP", "Platform Security Processor", "ARM-kärna i AMD-chipet som kör firmware innan OS startar"),
    ("PM table", "Power Management table", "Firmware-tabell med realtidsdata: temp, effekt, klockor"),
    ("MAC", "Multiply-Accumulate", "Ström-injicering i neuronen — vår primära insignal"),
    ("BVpar", "Breakdown Voltage (parallel)", "Spänningen där lavinström startar i memristorn (~0.55V)"),
    ("Vg", "Gate Voltage", "Styrsignal till memristorn — styr lavintröskel"),
    ("vmem", "Membrane Voltage", "Neuronens 'laddning' — ackumuleras tills den spikar"),
    ("LFSR", "Linear Feedback Shift Register", "Pseudoslumpgenerator i FPGA — driver lavinmodellen"),
    ("SOC", "Self-Organized Criticality", "System som naturligt hamnar på gränsen mellan ordning och kaos"),
    ("PSD", "Power Spectral Density", "Hur energi fördelas över frekvenser — 1/f = lutning -1"),
    ("1/f", "Ett-över-f (rosa brus)", "Brus där lågfrekventa variationer dominerar — finns i hjärnan och i båda våra substrat"),
    ("MC", "Memory Capacity", "Mått på hur mycket historik ett reservoir kan lagra"),
    ("R²", "Förklarad varians", "0 = ingen korrelation, 1 = perfekt prediktion"),
    ("EI", "Effective Information", "Mått på kausal kraft — hur mycket makronivån 'orsakar'"),
    ("ISI", "Inter-Spike Interval", "Tid mellan två spikar — dess variation avslöjar dynamiken"),
    ("CV", "Coefficient of Variation", "Standardavvikelse / medelvärde — mått på variabilitet"),
    ("TE", "Transfer Entropy", "Informationsflöde i bitar från en signal till en annan"),
    ("MI", "Mutual Information", "Delad information mellan två signaler (i bitar)"),
    ("UART", "Universal Async Receiver/Transmitter", "Seriell kommunikation — vår 8-neuron-brygga"),
    ("UDP", "User Datagram Protocol", "Nätverksprotokoll — vår 128-neuron Ethernet-brygga"),
    ("CRC", "Cyclic Redundancy Check", "Felkontroll — varje paket har CRC8"),
    ("Q16.16", "Fixed-point 16.16", "Talformat: 16 bitar heltal + 16 bitar decimal (istället för float)"),
    ("DTC", "DT over C (dt/C)", "Tidsstegets storlek delat med kapacitans — styr membranuppdatering"),
    ("τ", "Tau (tidskonstant)", "Hur snabbt membranet 'glömmer' — vårt τ=210ms (hjärnlikt)"),
    ("HfO₂", "Hafniumdioxid", "Materialet i memristorn — tunnt lager där filament bildas"),
    ("RTL", "Register Transfer Level", "Verilog-kod som beskriver FPGA-logik"),
    ("XDC", "Xilinx Design Constraints", "Fil som mappar RTL-signaler till fysiska pinnar"),
    ("hwmon", "Hardware Monitor", "Linux sysfs-gränssnitt för temp, effekt, fläkthastighet"),
    ("sysfs", "System Filesystem", "Virtuellt filsystem i Linux för att läsa/skriva hårdvara"),
]

table = doc.add_table(rows=1, cols=3)
table.style = 'Medium Shading 1 Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr = table.rows[0].cells
hdr[0].text = 'Akronym'
hdr[1].text = 'Står för'
hdr[2].text = 'Enkel förklaring'

for cell in hdr:
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(10)

for acronym, full, desc in rows:
    row = table.add_row().cells
    row[0].text = acronym
    row[1].text = full
    row[2].text = desc
    # Bold the acronym
    for p in row[0].paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(10)
    for p in row[1].paragraphs:
        for run in p.runs:
            run.font.size = Pt(10)
    for p in row[2].paragraphs:
        for run in p.runs:
            run.font.size = Pt(10)

# Set column widths
for row in table.rows:
    row.cells[0].width = Inches(1.0)
    row.cells[1].width = Inches(2.5)
    row.cells[2].width = Inches(3.5)

out = 'docs/FEEL_acronyms.docx'
import os
os.makedirs('docs', exist_ok=True)
doc.save(out)
print(f"Saved: {out}")
